import random
import pygame
from PIL import Image
from fpdf import FPDF
from docx import Document
from docx.shared import Inches


class Puzzle:
    def __init__(self, imagen, filas: int, columnas: int):
        self.filas = filas
        self.columnas = columnas
        self.imagen = imagen

    def cargar_imagen(self):
        ancho = self.imagen.get_width() // self.columnas # Obtenemos el ancho
        alto = self.imagen.get_height() // self.filas # Obtenemos el alto  // DIV en entero
        piezas = []

        for fil in range(self.filas):
            for col in range(self.columnas):
                x = col * ancho
                y = fil * alto
                rect = pygame.Rect(x, y, ancho, alto) #Generar el rectangulo en la interfaz
                pieza = self.imagen.subsurface(rect).copy() # creamos el recorte diferente al original
                piezas.append(pieza) # agregamos la pieza a la lista
        return piezas

    def mover_teclado(self, columnas: int, mover=None, piezas=None):
        pygame.init()
        pantalla = pygame.display.set_mode((400, 400))
        pygame.display.set_caption("Mover con teclado")
        reloj = pygame.time.Clock()
        ejecutando = True

        while ejecutando:
            for evento in pygame.event.get():
                if evento.type == pygame.QUIT:
                    ejecutando = False
                elif evento.type == pygame.KEYDOWN:
                    if evento.key == pygame.K_UP or evento.key == pygame.K_w:
                        mover("up")
                    elif evento.key == pygame.K_DOWN or evento.key == pygame.K_s:
                        mover("down")
                    elif evento.key == pygame.K_LEFT or evento.key == pygame.K_a:
                        mover("left")
                    elif evento.key == pygame.K_RIGHT or evento.key == pygame.K_d:
                        mover("right")

            pantalla.fill((30, 30, 30))
            font = pygame.font.Font(None, 36)
            for i, pieza in enumerate(piezas):
                x = (i % columnas) * 100
                y = (i // columnas) * 100
                if pieza != 0:
                    texto = font.render(str(pieza), True, (255, 255, 255))
                    pantalla.blit(texto, (x + 30, y + 30))
            pygame.display.flip()
            reloj.tick(30)

    def multijugador(self, crear_puzzle, mover=None, verificar_ganador=None, dibujar_puzzle=None):
        piezas1, vacio1 = crear_puzzle()
        piezas2, vacio2 = crear_puzzle()
        ganador = None

        ancho_total = 800
        alto_total = 400
        pantalla = pygame.display.set_mode((ancho_total, alto_total))
        pygame.display.set_caption("Puzzle Multijugador")
        reloj = pygame.time.Clock()
        ejecutando = True

        while ejecutando:
            for evento in pygame.event.get():
                if evento.type == pygame.QUIT:
                    ejecutando = False
                elif evento.type == pygame.KEYDOWN:
                    if evento.key == pygame.K_UP:
                        vacio1 = mover(piezas1, vacio1, "up")
                    elif evento.key == pygame.K_DOWN:
                        vacio1 = mover(piezas1, vacio1, "down")
                    elif evento.key == pygame.K_LEFT:
                        vacio1 = mover(piezas1, vacio1, "left")
                    elif evento.key == pygame.K_RIGHT:
                        vacio1 = mover(piezas1, vacio1, "right")
                    elif evento.key == pygame.K_w:
                        vacio2 = mover(piezas2, vacio2, "up")
                    elif evento.key == pygame.K_s:
                        vacio2 = mover(piezas2, vacio2, "down")
                    elif evento.key == pygame.K_a:
                        vacio2 = mover(piezas2, vacio2, "left")
                    elif evento.key == pygame.K_d:
                        vacio2 = mover(piezas2, vacio2, "right")

            pantalla.fill((30, 30, 30))
            dibujar_puzzle(pantalla, piezas1, x_offset=0)
            dibujar_puzzle(pantalla, piezas2, x_offset=ancho_total // 2)
            pygame.draw.line(pantalla, (255, 255, 255), (ancho_total // 2, 0), (ancho_total // 2, alto_total), 4)
            pygame.display.flip()
            reloj.tick(30)

            if not ganador:
                if verificar_ganador(piezas1):
                    ganador = "¡Jugador 1 (Flechas) gana!"
                elif verificar_ganador(piezas2):
                    ganador = "¡Jugador 2 (WASD) gana!"
                if ganador:
                    return ganador


class TiposDeImagenes:
    def __init__(self, JPG: bool, PNG: bool, PDF: bool, WORD: bool):
        self.JPG = JPG
        self.PNG = PNG
        self.PDF = PDF
        self.WORD = WORD

    def guardar_imagen(self, imagen, nombre="Puzzle_done"):
        imagen_rgb = imagen.convert("RGB")

        if self.JPG:
            imagen_rgb.save(f"{nombre}.jpg")
        if self.PNG:
            imagen.save(f"{nombre}.png")
        if self.PDF:
            pdf = FPDF()
            pdf.add_page()
            pdf.image(f"{nombre}.jpg", 10, 10, 210, 297)
            pdf.output(f"{nombre}.pdf")
        if self.WORD:
            doc = Document()
            doc.add_heading(nombre, 0)
            doc.add_picture(f"{nombre}.jpg", width=Inches(5))
            doc.save(f"{nombre}.docx")

